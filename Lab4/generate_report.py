#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate Lab4 OO Analysis & Design Report (Word format)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

UML_DIR = "/Users/liqiulin/01_Projects/软件工程26春/Lab4/uml"
OUTPUT = "/Users/liqiulin/01_Projects/软件工程26春/Lab4/Lab4-OO分析与设计-实验报告.docx"

# All strings with Chinese quotes use LQ/RQ variables to avoid syntax issues
LQ = "\u201c"  # left Chinese quote "
RQ = "\u201d"  # right Chinese quote "


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_with_header(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.name = "宋体"
        set_cell_shading(cell, "D9E2F3")
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = "宋体"
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def set_heading_style(heading, font_name="黑体", font_size=None):
    for run in heading.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        if font_size:
            run.font.size = font_size


def add_centered_image(doc, img_path, width_inches=5.5, caption=""):
    doc.add_picture(img_path, width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5

    # ==================== Title Page ====================
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026年春季学期")
    run.font.size = Pt(16)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("计算学部《软件工程》课程")
    run.font.size = Pt(16)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("实验报告")
    run.font.size = Pt(26)
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Lab4 OO分析与设计")
    run.font.size = Pt(22)
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    for _ in range(4):
        doc.add_paragraph()

    info_table = doc.add_table(rows=4, cols=3)
    info_table.style = "Table Grid"
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("姓名", "", ""),
        ("学号", "", ""),
        ("联系方式", "", ""),
        ("", "Email/手机号码", "Email/手机号码"),
    ]
    for r, row_data in enumerate(info_data):
        for c, val in enumerate(row_data):
            cell = info_table.rows[r].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(12)

    doc.add_page_break()

    # ==================== TOC ====================
    h = doc.add_heading("目  录", level=1)
    set_heading_style(h)
    doc.add_paragraph("（文档全部完成之后，请更新目录）")
    doc.add_page_break()

    # ==================== Chapter 1 ====================
    h = doc.add_heading("1  本组项目概述", level=1)
    set_heading_style(h)

    h2 = doc.add_heading("1.1  业务目标与业务价值", level=2)
    set_heading_style(h2)
    doc.add_paragraph(
        "Aura Health 是一款基于大模型技术的办公室健康监测系统。其业务目标是为办公室人群提供全方位的健康状态追踪与智能干预，"
        "帮助企业和个人降低因久坐、缺乏运动、作息不规律等引发的职业健康风险。系统通过实时采集心率、血压、睡眠、饮水、步数等多维生理指标，"
        "结合 AI 大模型生成个性化健康建议与异常预警，实现从" + LQ + "被动记录" + RQ + "到" + LQ + "主动干预" + RQ + "的健康管理闭环。"
    )
    doc.add_paragraph(
        "业务价值体现在：（1）提升员工健康意识与生活质量，降低因病缺勤率；（2）为企业 HR/行政提供员工健康画像与群体分析能力；"
        "（3）积累健康数据资产，支撑长期健康趋势分析与保险精算。"
    )

    h2 = doc.add_heading("1.2  功能概述", level=2)
    set_heading_style(h2)
    doc.add_paragraph("系统采用 B/S 架构，前后端分离，核心功能包括：")

    features = [
        "用户管理：注册、登录、个人信息维护，JWT 鉴权机制保障安全",
        "健康数据采集与展示：手动录入心率、血压、体重、睡眠、饮水、步数；仪表盘实时数据卡片 + 14天趋势折线图",
        "AI 健康助手：基于 SiliconFlow 大模型（Qwen/DeepSeek），结合用户健康数据生成个性化建议；对话式交互，支持连续上下文",
        "智能提醒：久坐提醒、饮水提醒、异常预警（心率/血压超标时触发）",
        "数据报告：AI 生成周报/月报，支持 PDF 健康报告导出",
    ]
    for f in features:
        doc.add_paragraph(f, style="List Bullet")

    h2 = doc.add_heading("1.3  技术栈", level=2)
    set_heading_style(h2)
    add_table_with_header(doc, ["层级", "技术", "说明"], [
        ["前端", "React 18 + Vite + Ant Design", "现代化 SPA 单页应用"],
        ["后端", "Python FastAPI + SQLAlchemy", "高性能异步 API 框架"],
        ["数据库", "SQLite（原型）/ PostgreSQL（生产）", "关系型数据库"],
        ["AI/LLM", "SiliconFlow API (Qwen/DeepSeek)", "大语言模型推理服务"],
        ["可视化", "Recharts", "健康数据趋势图表"],
    ])

    doc.add_page_break()

    # ==================== Chapter 2 ====================
    h = doc.add_heading("2  用例模型", level=1)
    set_heading_style(h)

    h2 = doc.add_heading("2.1  角色清单", level=2)
    set_heading_style(h2)
    doc.add_paragraph("本系统的 Actor 如下：")
    add_table_with_header(doc, ["Actor", "类型", "说明"], [
        ["普通用户", "人员（主要 Actor）", "注册登录、录入/查看健康数据、AI咨询、接收提醒、导出报告"],
        ["AI大模型服务", "外部软件系统", "SiliconFlow API，提供智能健康建议生成能力"],
        ["系统时钟", "系统时钟", "定时触发久坐提醒、饮水提醒、异常预警"],
    ])

    h2 = doc.add_heading("2.2  用例模型", level=2)
    set_heading_style(h2)
    doc.add_paragraph("系统用例模型如下图所示，包含3个 Actor 和12个用例，用例之间存在 include 和 extend 关系：")
    add_centered_image(doc, os.path.join(UML_DIR, "UseCaseDiagram.png"), 5.8, "图 2-1 用例模型图")

    doc.add_paragraph("用例清单说明：")
    add_table_with_header(doc, ["用例编号", "用例名称", "参与 Actor", "说明"], [
        ["UC1", "用户注册", "普通用户", "填写用户名、密码等个人信息创建账号"],
        ["UC2", "用户登录", "普通用户", "输入用户名密码登录系统，获取JWT令牌"],
        ["UC3", "健康数据录入", "普通用户", "录入心率、血压、体重、睡眠、饮水、步数"],
        ["UC4", "查看健康仪表盘", "普通用户", "查看最新健康数据概览和核心生命体征"],
        ["UC5", "查看历史趋势", "普通用户", "查看14天健康数据趋势折线图"],
        ["UC6", "AI健康咨询", "普通用户、AI大模型服务", "与AI助手对话，获取个性化健康建议"],
        ["UC7", "管理个人信息", "普通用户", "查看和修改个人基本信息"],
        ["UC8", "接收智能提醒", "普通用户、系统时钟", "接收久坐、饮水、异常预警提醒"],
        ["UC8a", "久坐提醒", "系统时钟", "根据活动量数据触发久坐提醒"],
        ["UC8b", "饮水提醒", "系统时钟", "定时触发饮水提醒"],
        ["UC8c", "异常预警", "系统时钟", "心率/血压超出正常范围时触发预警"],
        ["UC9", "导出健康报告", "普通用户", "生成并下载PDF健康报告"],
    ])

    # Use case 1
    h2 = doc.add_heading("2.3  用例1的事件流描述：用户注册与登录", level=2)
    set_heading_style(h2)
    doc.add_paragraph("本用例所涉及的 Actor(s)：普通用户")
    p = doc.add_paragraph()
    run = p.add_run("常规事件流：")
    run.bold = True
    flows = [
        "1. 用户在登录页面选择" + LQ + "注册" + RQ + "，进入注册表单",
        "2. 用户输入用户名、密码、真实姓名、年龄、职业等个人信息",
        "3. 系统验证用户名是否已存在",
        "4. 系统对密码进行哈希加密（sha256_crypt）后存储用户信息到数据库",
        "5. 系统返回注册成功提示，引导用户前往登录",
        "6. 用户在登录页面输入用户名和密码",
        "7. 系统根据用户名查询数据库验证用户身份",
        "8. 系统对输入密码进行哈希比对验证",
        "9. 验证通过，系统生成 JWT 令牌（有效期24小时）返回给客户端",
        "10. 客户端将令牌存储至 localStorage，跳转至仪表盘主页",
    ]
    for f in flows:
        doc.add_paragraph(f)

    p = doc.add_paragraph()
    run = p.add_run("备选事件流：")
    run.bold = True
    alt_flows = [
        "3a. 用户名已存在：系统提示" + LQ + "用户名已存在" + RQ + "，要求更换用户名重新输入",
        "8a. 密码不匹配：系统提示" + LQ + "用户名或密码错误" + RQ,
        "9a. JWT令牌过期后：客户端检测到401响应，自动跳转至登录页",
    ]
    for f in alt_flows:
        doc.add_paragraph(f)

    # Use case 2
    h2 = doc.add_heading("2.4  用例2的事件流描述：健康数据录入", level=2)
    set_heading_style(h2)
    doc.add_paragraph("本用例所涉及的 Actor(s)：普通用户")
    p = doc.add_paragraph()
    run = p.add_run("常规事件流：")
    run.bold = True
    flows = [
        "1. 用户在侧边导航栏点击" + LQ + "数据录入" + RQ + "进入录入页面",
        "2. 系统展示数据录入表单，分为" + LQ + "生命体征监测" + RQ + "（心率、收缩压、舒张压）和" + LQ + "生活质量指数" + RQ + "（体重、睡眠、饮水、步数）两个区域",
        "3. 用户在" + LQ + "生命体征监测" + RQ + "区域填写心率（BPM）、收缩压（mmHg）、舒张压（mmHg）",
        "4. 用户在" + LQ + "生活质量指数" + RQ + "区域填写体重（kg）、睡眠时长（hrs）、饮水量（ml）、步数（steps）",
        "5. 用户点击" + LQ + "同步至 Aura 全球库" + RQ + "按钮提交数据",
        "6. 系统对输入数据进行范围校验（心率30-220、血压60-250/40-150、体重30-300、睡眠0-24、饮水0-10000、步数0-100000）",
        "7. 校验通过，系统将数据与当前登录用户关联，存入数据库并标记记录日期",
        "8. 系统检查录入数据是否存在异常指标（如收缩压>140或舒张压>90）",
        "9. 若存在异常，系统创建异常预警提醒并关联至当前用户",
        "10. 系统返回" + LQ + "健康数据已同步" + RQ + "成功提示，仪表盘数据自动刷新",
    ]
    for f in flows:
        doc.add_paragraph(f)

    p = doc.add_paragraph()
    run = p.add_run("备选事件流：")
    run.bold = True
    alt_flows = [
        "6a. 心率超出合理范围（<30或>220）：系统提示" + LQ + "心率数据不合理，请检查" + RQ,
        "6b. 血压数据超出范围：系统提示具体字段的校验错误信息",
        "6c. 部分字段为空：系统允许部分数据为空（默认值为0），仅保存已填写项",
        "7a. 网络异常或服务不可用：系统提示" + LQ + "同步失败" + RQ + "，建议稍后重试",
    ]
    for f in alt_flows:
        doc.add_paragraph(f)

    # Use case 3
    h2 = doc.add_heading("2.5  用例3的事件流描述：AI健康咨询", level=2)
    set_heading_style(h2)
    doc.add_paragraph("本用例所涉及的 Actor(s)：普通用户、AI大模型服务")
    p = doc.add_paragraph()
    run = p.add_run("常规事件流：")
    run.bold = True
    flows = [
        "1. 用户点击侧边导航栏" + LQ + "AI助手" + RQ + "进入对话页面",
        "2. 系统从数据库加载并展示用户的历史对话记录",
        "3. 用户在底部输入框输入健康相关问题（如" + LQ + "我最近血压偏高怎么办？" + RQ + "）",
        "4. 系统将用户消息以 role=user 存入 ai_conversations 表",
        "5. 系统从数据库获取用户最近7天的健康记录作为健康数据上下文",
        "6. 系统获取最近10条对话历史（按时间正序排列）",
        "7. 系统构建包含系统提示词 + 用户个人信息 + 健康数据 + 对话历史的完整消息列表",
        "8. 系统通过 HTTPS 调用外部 AI 大模型 API（SiliconFlow），发送消息列表",
        "9. AI大模型返回基于上下文的个性化健康建议",
        "10. 系统将AI回复以 role=assistant 存入 ai_conversations 表",
        "11. 系统在对话界面实时展示AI回复，对话界面自动滚动至最新消息",
    ]
    for f in flows:
        doc.add_paragraph(f)

    p = doc.add_paragraph()
    run = p.add_run("备选事件流：")
    run.bold = True
    alt_flows = [
        "8a. AI大模型API不可用或请求超时（15秒）：系统自动切换至本地规则引擎，基于关键词匹配生成兜底回复",
        "8b. AI大模型返回异常响应：系统提示" + LQ + "远端模型故障" + RQ + "并附带本地规则引擎的回复",
        "2a. 无历史对话记录：系统展示欢迎引导界面，提示用户可以先录入健康数据以获得更精准的建议",
        "5a. 无健康记录：上下文中标注" + LQ + "暂无健康记录数据" + RQ + "，AI仍可进行通用健康问答",
    ]
    for f in alt_flows:
        doc.add_paragraph(f)

    doc.add_page_break()

    # ==================== Chapter 3 ====================
    h = doc.add_heading("3  类识别", level=1)
    set_heading_style(h)

    h2 = doc.add_heading("3.1  边界类", level=2)
    set_heading_style(h2)
    add_table_with_header(doc, ["类名（中文）", "类名（英文）", "类的作用概述"], [
        ["登录界面", "LoginView", "用户登录/注册的交互界面，收集用户名密码并展示认证结果"],
        ["仪表盘界面", "DashboardView", "展示用户最新健康数据概览，包括核心生命体征、每日动态、趋势图表等"],
        ["数据录入界面", "DataEntryView", "提供健康数据录入表单，收集心率、血压、体重、睡眠、饮水、步数等数据"],
        ["AI对话界面", "AIChatView", "AI健康咨询的对话交互界面，展示历史消息和输入框，支持实时对话"],
        ["个人信息界面", "ProfileView", "用户查看和修改个人基本信息（姓名、年龄、性别、身高、体重、职业）"],
        ["提醒通知界面", "ReminderNotificationView", "展示久坐、饮水、异常预警等智能提醒通知弹窗"],
        ["报告导出界面", "ReportExportView", "PDF健康报告的预览、生成与下载界面"],
    ], col_widths=[3, 4, 8])

    h2 = doc.add_heading("3.2  控制类", level=2)
    set_heading_style(h2)
    add_table_with_header(doc, ["类名（中文）", "类名（英文）", "类的作用概述"], [
        ["用户认证控制器", "AuthController", "处理用户注册、登录验证、JWT令牌生成与验证、密码哈希加密"],
        ["健康数据控制器", "HealthDataController", "处理健康数据的增删查改，数据范围校验，数据排序与过滤"],
        ["AI咨询控制器", "AIConsultController", "处理AI对话消息的发送与接收，构建提示词上下文，管理对话历史，包含本地规则引擎兜底逻辑"],
        ["仪表盘控制器", "DashboardController", "聚合用户最新健康数据和近期趋势，为仪表盘界面提供展示数据"],
        ["提醒控制器", "ReminderController", "根据系统时钟和健康数据触发久坐/饮水/异常提醒，管理提醒的创建和已读状态"],
        ["报告控制器", "ReportController", "生成PDF健康报告，调用AI生成周报/月报总结，处理报告下载请求"],
    ], col_widths=[3.5, 4.5, 7])

    h2 = doc.add_heading("3.3  实体类", level=2)
    set_heading_style(h2)
    add_table_with_header(doc, ["类名（中文）", "类名（英文）", "类的作用概述"], [
        ["用户", "User", "存储用户基本信息（用户名、密码哈希、姓名、年龄、性别、身高、体重、职业、创建时间）"],
        ["健康记录", "HealthRecord", "存储用户的健康数据（心率、血压、体重、睡眠、饮水、步数）及记录日期、创建时间"],
        ["AI对话", "AIConversation", "存储用户与AI助手的对话消息（角色user/assistant、内容、创建时间）"],
        ["提醒", "Reminder", "存储提醒记录（类型sedentary/water/abnormal、消息内容、触发时间、是否已读）"],
    ], col_widths=[3, 4, 8])

    doc.add_page_break()

    # ==================== Chapter 4 ====================
    h = doc.add_heading("4  领域模型", level=1)
    set_heading_style(h)

    h2 = doc.add_heading("4.1  分析类图", level=2)
    set_heading_style(h2)
    doc.add_paragraph(
        "分析类图描述边界类、控制类、实体类之间的逻辑关联关系。本图中无需给出各类的属性和方法，"
        "仅展示三类分析类之间的调用与依赖关系。图中所有类均已在第3部分的三个表格中定义，类名使用中文。"
    )
    add_centered_image(doc, os.path.join(UML_DIR, "AnalysisClassDiagram.png"), 5.8, "图 4-1 分析类图（边界类-控制类-实体类 关联关系）")

    doc.add_paragraph("分析类图核心关联说明：")
    relations = [
        "LoginView -> AuthController -> User：登录界面通过认证控制器访问用户实体",
        "DashboardView -> DashboardController -> User, HealthRecord：仪表盘通过控制器聚合用户信息和健康记录",
        "DataEntryView -> HealthDataController -> User, HealthRecord：数据录入界面通过控制器创建健康记录",
        "HealthDataController -> ReminderController：数据录入时若检测异常，触发提醒控制器",
        "AIChatView -> AIConsultController -> User, AIConversation, HealthRecord：AI对话需要用户信息、对话历史和健康记录上下文",
        "ReminderNotificationView -> ReminderController -> User, Reminder, HealthRecord：提醒通知依赖用户、提醒记录和健康数据",
        "ProfileView -> AuthController -> User：个人信息管理通过认证控制器操作用户实体",
        "ReportExportView -> ReportController -> User, HealthRecord, AIConversation：报告导出需要整合用户信息、健康数据和AI分析",
    ]
    for r in relations:
        doc.add_paragraph(r, style="List Bullet")

    h2 = doc.add_heading("4.2  实体类图（领域类图）", level=2)
    set_heading_style(h2)
    doc.add_paragraph(
        "领域类图刻画各实体类的主要属性集合（可见性、名称、数据类型）、主要操作集合（可见性、名称、返回值数据类型、参数列表）、"
        "类之间的关系（组合、依赖），以及关系的角色名、多重性、方向等信息。图中所有实体类均已在第3.3节的表格中定义，"
        "类名、属性名、方法名、参数名均使用中文。"
    )
    add_centered_image(doc, os.path.join(UML_DIR, "DomainClassDiagram.png"), 5.5, "图 4-2 领域类图（实体类图）")

    doc.add_paragraph("实体类间关系说明：")
    doc.add_paragraph(
        "User -- 0..* HealthRecord（组合关系）：一个用户拥有零到多条健康记录，健康记录生命周期依赖于用户，"
        "用户删除时其所有健康记录一同删除。角色名" + LQ + "拥有" + RQ + "，多重性 1 对 0..*。"
    )
    doc.add_paragraph(
        "User -- 0..* AIConversation（组合关系）：一个用户拥有零到多条AI对话消息，对话消息生命周期依赖于用户。"
        "角色名" + LQ + "拥有" + RQ + "，多重性 1 对 0..*。"
    )
    doc.add_paragraph(
        "User -- 0..* Reminder（组合关系）：一个用户拥有零到多条提醒记录，提醒生命周期依赖于用户。"
        "角色名" + LQ + "拥有" + RQ + "，多重性 1 对 0..*。"
    )
    doc.add_paragraph(
        "HealthRecord -- Reminder（依赖关系）：异常健康记录可能触发提醒创建，属于依赖关系。"
        "角色名" + LQ + "触发" + RQ + "，多重性 0..1 对 0..*。"
    )

    doc.add_page_break()

    # ==================== Chapter 5 ====================
    h = doc.add_heading("5  时序模型", level=1)
    set_heading_style(h)

    h2 = doc.add_heading("5.1  用例1的时序模型：用户注册与登录", level=2)
    set_heading_style(h2)
    doc.add_paragraph(
        "该时序模型描述用户注册与登录过程中，LoginView（边界类）、AuthController（控制类）、User（实体类）之间的消息流和调用关系。"
        "所有对象已在第3部分被识别，操作已在4.2节的类图中体现。"
    )
    add_centered_image(doc, os.path.join(UML_DIR, "Sequence_RegisterLogin.png"), 5.5, "图 5-1 用户注册与登录时序图")

    doc.add_paragraph("时序说明：")
    seq1 = [
        "注册阶段：用户通过LoginView输入注册信息 -> AuthController查询User实体验证用户名唯一性 -> 哈希加密密码 -> 保存用户信息 -> 返回注册结果",
        "登录阶段：用户通过LoginView输入登录信息 -> AuthController查询User实体 -> 验证密码（哈希比对） -> 生成JWT令牌 -> 返回给客户端存储",
    ]
    for s in seq1:
        doc.add_paragraph(s, style="List Bullet")

    h2 = doc.add_heading("5.2  用例2的时序模型：健康数据录入", level=2)
    set_heading_style(h2)
    doc.add_paragraph(
        "该时序模型描述健康数据录入过程中，DataEntryView（边界类）、HealthDataController（控制类）、HealthRecord（实体类）、"
        "ReminderController（控制类）、Reminder（实体类）之间的消息流和调用关系。"
    )
    add_centered_image(doc, os.path.join(UML_DIR, "Sequence_DataEntry.png"), 5.5, "图 5-2 健康数据录入时序图")

    doc.add_paragraph("时序说明：")
    seq2 = [
        "正常流程：用户填写数据 -> DataEntryView提交 -> HealthDataController校验 -> 创建HealthRecord -> 检查异常 -> 若异常则触发ReminderController创建Reminder",
        "异常分支：数据校验失败时直接返回错误提示，不进入存储流程",
    ]
    for s in seq2:
        doc.add_paragraph(s, style="List Bullet")

    h2 = doc.add_heading("5.3  用例3的时序模型：AI健康咨询", level=2)
    set_heading_style(h2)
    doc.add_paragraph(
        "该时序模型描述AI健康咨询过程中，AIChatView（边界类）、AIConsultController（控制类）、AIConversation（实体类）、"
        "HealthRecord（实体类）、AI大模型服务（外部系统）之间的消息流和调用关系。"
    )
    add_centered_image(doc, os.path.join(UML_DIR, "Sequence_AIChat.png"), 5.5, "图 5-3 AI健康咨询时序图")

    doc.add_paragraph("时序说明：")
    seq3 = [
        "正常流程：用户发送消息 -> AIConsultController保存用户消息 -> 获取HealthRecord上下文 -> 获取AIConversation历史 -> 构建提示词 -> 调用AI大模型API -> 保存AI回复 -> 展示",
        "兜底流程：AI大模型不可用时，切换至本地规则引擎，基于关键词匹配生成回复，保证服务可用性",
    ]
    for s in seq3:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_page_break()

    # ==================== Chapter 6 ====================
    h = doc.add_heading("6  部署模型", level=1)
    set_heading_style(h)

    doc.add_paragraph(
        "系统采用 B/S（Browser/Server）架构，用户通过 PC 或手机浏览器访问系统。"
        "后端应用服务器部署 FastAPI 服务和数据库，通过 HTTPS 与前端通信；"
        "AI 大模型能力通过外部 SiliconFlow 云端 API 以 HTTPS 方式调用。"
    )
    add_centered_image(doc, os.path.join(UML_DIR, "DeploymentDiagram.png"), 5.5, "图 6-1 部署模型图")

    doc.add_paragraph("部署节点说明：")
    add_table_with_header(doc, ["节点", "部署内容", "说明"], [
        ["客户端节点", "React SPA (LoginView, DashboardView, DataEntryView, AIChatView)", "用户通过 PC/手机浏览器访问，前端单页应用"],
        ["应用服务器节点", "FastAPI 后端服务 + SQLite/PostgreSQL 数据库", "部署在云服务器或 Docker 容器中，处理业务逻辑和数据存储"],
        ["外部AI服务节点", "SiliconFlow 大模型推理API (Qwen/DeepSeek)", "云端第三方服务，通过 HTTPS API 调用"],
    ])

    doc.add_paragraph("网络连接说明：")
    net = [
        "客户端 -> 应用服务器：HTTPS 加密连接，前端通过 REST API 与后端通信",
        "应用服务器 -> 外部AI服务：HTTPS 加密连接，后端通过 OpenAI 兼容 API 格式调用大模型",
        "FastAPI 后端 -> 数据库：本地连接（SQLite）或内网连接（PostgreSQL）",
    ]
    for n in net:
        doc.add_paragraph(n, style="List Bullet")

    doc.add_page_break()

    # ==================== Chapter 7 ====================
    h = doc.add_heading("7  小结", level=1)
    set_heading_style(h)

    h2 = doc.add_heading("7.1  建模是否有必要？", level=2)
    set_heading_style(h2)
    doc.add_paragraph(
        "建模是非常有必要的。直接写代码虽然看起来节省了前期的建模时间，但往往会导致后期大量返工。"
        "通过本次实验的建模实践，我们深刻体会到：建模迫使我们在编码前系统性地思考系统的整体架构、模块间的交互逻辑和数据的流转方式。"
        "在绘制用例图时，我们发现了一些之前需求分析中遗漏的边界场景（如AI服务不可用时的兜底逻辑）；"
        "在绘制类图时，我们明确了每个类的职责边界，避免了编码时可能出现的职责混乱。"
        "边写代码边思考的方式容易陷入" + LQ + "补丁式" + RQ + "开发——发现问题就局部修补，最终导致系统结构混乱、代码耦合度高。"
        "建模提供了一个全局视角，让我们在动手写代码之前就对系统有了清晰、完整的认识。"
    )

    h2 = doc.add_heading("7.2  如何逼着自己建模时想得更细节？", level=2)
    set_heading_style(h2)
    doc.add_paragraph("我们总结出以下方法来确保建模时的思考深度：")
    methods = [
        "为每个用例编写完整的事件流描述，不仅写常规流程，还要穷举所有可能的备选事件流和异常情况。这一步强迫我们思考各种边界条件。",
        "在识别类的方法时，不仅要列出方法签名，还要设计其内部业务逻辑（如校验规则、计算公式、调用链路），把方法的输入-处理-输出想清楚。",
        "绘制时序图时，逐条追踪每一个消息的来源和去向，确保不存在" + LQ + "悬空" + RQ + "的消息——每条消息都必须对应某个类的某个操作。这一步能发现类图中的方法遗漏。",
        "从用例图到类图到时序图进行交叉验证：用例中的每个步骤都应在时序图中有对应消息，时序图中的每个操作都应在类图的方法中体现。",
        "角色扮演法：假想自己是系统中的某个对象，模拟一条用户请求从进入到返回的完整旅程，看是否每个环节都能被当前的模型覆盖。",
    ]
    for m in methods:
        doc.add_paragraph(m, style="List Number")

    h2 = doc.add_heading("7.3  对面向对象分析与设计方法的看法", level=2)
    set_heading_style(h2)
    doc.add_paragraph(
        "面向对象分析与设计方法的核心优势在于其与人类认知方式的高度契合——我们天然倾向于将世界理解为由各种" + LQ + "对象" + RQ + "及其交互构成。"
        "具体而言："
    )
    doc.add_paragraph(
        "优势：（1）封装性使模块之间松耦合，修改一个类的内部实现不会影响其他类，提升了系统的可维护性；"
        "（2）继承机制减少了重复代码，通过抽象父类统一管理公共属性和行为；"
        "（3）多态支持系统扩展，新增功能只需添加子类或实现接口，无需修改已有代码（开闭原则）；"
        "（4）面向对象的思维方式——识别实体、划清职责、定义交互——天然适合团队协作，每个成员可以独立负责一组类的开发。"
    )
    doc.add_paragraph(
        "劣势：（1）对于简单的、流程驱动的计算任务（如数据格式转换、批量处理），面向对象可能引入不必要的抽象层次，"
        "而结构化的" + LQ + "算法+数据结构" + RQ + "方式更为直接高效；"
        "（2）过度设计风险——初学者容易为假设的未来需求创建过多的抽象层，反而增加系统复杂度；"
        "（3）面向对象的设计质量高度依赖于设计者的经验，错误的抽象（如将本应是方法的操作抽象为类）会导致系统更加难以理解和维护。"
    )

    h2 = doc.add_heading("7.4  对UML建模语言的看法", level=2)
    set_heading_style(h2)
    doc.add_paragraph("UML 作为一种标准化的建模语言，提供了丰富的图表类型，各图表在我们的实践中发挥了不同的作用：")
    add_table_with_header(doc, ["图表类型", "作用", "我们的评价"], [
        ["用例图", "界定系统边界和功能范围，明确Actor与系统的交互", "非常实用，帮助团队统一对系统功能的认知，避免需求遗漏"],
        ["分析类图（鲁棒图）", "将系统要素分为边界/控制/实体三类，明确职责划分", "对初学者非常友好，强制思考UI-逻辑-数据的分离，但可能过于简化"],
        ["领域类图", "详细描述实体类的属性、方法和关系", "最有价值的图表，直接指导编码实现，但绘制成本较高"],
        ["时序图", "描述对象间的消息交互顺序", "对验证设计可行性极其重要，能发现类图中的方法遗漏和调用逻辑错误"],
        ["部署图", "描述系统的物理部署架构", "对运维和部署有指导意义，但对小型系统价值有限"],
    ])
    doc.add_paragraph(
        "总体而言，UML 的各图表确实起到了帮助开发者的作用，但并非每个项目都需要绘制所有类型的图表。"
        "我们认为应根据项目规模和复杂度选择性地使用：用例图和类图是必选项，时序图对核心用例必不可少，"
        "部署图在涉及多节点部署时才有显著价值。避免为画图而画图，建模的目的是辅助思考而非形式主义。"
    )

    h2 = doc.add_heading("7.5  是否愿意先建模再开发？", level=2)
    set_heading_style(h2)
    doc.add_paragraph(
        "我们愿意在以后的开发中采用" + LQ + "先建模、团队讨论修改、再开发" + RQ + "的流程。理由如下："
    )
    reasons = [
        "建模阶段的修改成本远低于编码阶段。在模型上修改一个类的关系可能只需要几分钟，但代码重构可能需要数小时甚至数天。",
        "团队围绕模型讨论可以尽早统一认知，避免因理解不一致导致的返工。用例图和类图为讨论提供了共同的" + LQ + "语言" + RQ + "。",
        "模型作为文档资产，在项目维护和人员更替时具有重要价值，新成员可以通过模型快速理解系统架构。",
        "当然，建模不宜过度细化到伪代码级别——建模的粒度应停留在" + LQ + "设计决策" + RQ + "层面，具体的实现细节留给编码阶段解决。",
    ]
    for r in reasons:
        doc.add_paragraph(r, style="List Number")

    # Save
    doc.save(OUTPUT)
    print("Report saved to: " + OUTPUT)


if __name__ == "__main__":
    main()